import streamlit as st
from docx import Document
import smtplib
from email.message import EmailMessage
import os
import tempfile
from datetime import date

st.set_page_config(page_title="Pré-Avaliação Neuropsicológica - Adulto", page_icon="🧠")
st.title("Formulário de Pré-Avaliação Neuropsicológica - Adulto")

st.markdown("**Finalidade:** Este questionário tem por objetivo levantar dados sobre o desenvolvimento, histórico médico, desempenho acadêmico e profissional do(a) avaliado(a), a fim de subsidiar a análise neuropsicológica.")

# ==== DADOS BÁSICOS ====
st.header("Dados de Identificação")
nome = st.text_input("Nome completo")
data_avaliacao = st.date_input("Data da Avaliação", value=date.today())
tipo_avaliado = st.radio("Tipo de respondente", ["Paciente", "Outro"], index=0)
grau_parentesco = st.text_input("Grau de parentesco") if tipo_avaliado == "Outro" else ""

telefone = st.text_input("Telefone")
data_nascimento = st.date_input("Data de nascimento")
idade = st.number_input("Idade", min_value=0, max_value=120)
sexo = st.text_input("Sexo")
endereco = st.text_input("Endereço completo")
cidade_estado = st.text_input("Cidade/Estado de nascimento")
mao_dominante = st.radio("Mão dominante", ["Direita", "Esquerda"])
idiomas = st.text_input("Quais idiomas?") if st.checkbox("Fala outro idioma?") else ""

diagnosticos = st.text_area("Diagnóstico(s) médico(s)")
encaminhado_por = st.text_input("Encaminhado por")
data_acidente = st.date_input("Data do acidente/início da doença (se aplicável)")
cuidador_grau = st.text_input("Grau de parentesco do cuidador") if st.checkbox("Possui cuidador?") else ""

# ==== FUNCIONAMENTO ====
st.header("Seção 1: Funcionamento Físico")
sintomas_fisicos = st.multiselect("Sintomas físicos", [...])
outros_fisico = st.text_input("Outros sintomas físicos")

st.header("Seção 2: Funcionamento Sensorial")
sensoriais = st.multiselect("Sintomas sensoriais", [...])
outros_sens = st.text_input("Outros sintomas sensoriais")

st.header("Seção 3: Cognição")
cognicao = st.multiselect("Dificuldades cognitivas", [...])
obs_cognicao = st.text_area("Observações cognitivas")

st.header("Seção 4: Linguagem e Matemática")
linguagem = st.multiselect("Dificuldades em linguagem/matemática", [...])
outros_ling = st.text_input("Outros aspectos de linguagem")

st.header("Seção 5: Habilidades Não Verbais")
nverbais = st.multiselect("Habilidades não verbais", [...])
outros_nv = st.text_input("Outros aspectos não verbais")

st.warning("Este documento é um modelo automatizado e deve ser revisado pelo(a) psicólogo(a) responsável antes de qualquer uso clínico.")

# ==== ENVIO ====
if st.button("Enviar"):
    if not nome or not telefone:
        st.error("Preencha ao menos o nome e o telefone para prosseguir.")
    else:
        doc = Document()
        doc.add_heading("Formulário de Pré-Avaliação Neuropsicológica - Adulto", 0)
        doc.add_paragraph(f"Nome: {nome}")
        doc.add_paragraph(f"Data: {data_avaliacao}")
        doc.add_paragraph(f"Tipo de respondente: {tipo_avaliado} {grau_parentesco}")
        doc.add_paragraph(f"Telefone: {telefone}, Nascimento: {data_nascimento}, Idade: {idade}, Sexo: {sexo}")
        doc.add_paragraph(f"Endereço: {endereco}, Cidade/Estado: {cidade_estado}, Mão dominante: {mao_dominante}")
        doc.add_paragraph(f"Idiomas: {idiomas}")
        doc.add_paragraph(f"Diagnóstico(s): {diagnosticos}, Encaminhado por: {encaminhado_por}, Acidente: {data_acidente}")
        doc.add_paragraph(f"Cuidador: {cuidador_grau}")
        doc.add_heading("Funcionamento Físico", level=1)
        doc.add_paragraph(", ".join(sintomas_fisicos) + ". Outros: " + outros_fisico)
        doc.add_heading("Funcionamento Sensorial", level=1)
        doc.add_paragraph(", ".join(sensoriais) + ". Outros: " + outros_sens)
        doc.add_heading("Cognição", level=1)
        doc.add_paragraph(", ".join(cognicao) + ". Observações: " + obs_cognicao)
        doc.add_heading("Linguagem e Matemática", level=1)
        doc.add_paragraph(", ".join(linguagem) + ". Outros: " + outros_ling)
        doc.add_heading("Habilidades Não Verbais", level=1)
        doc.add_paragraph(", ".join(nverbais) + ". Outros: " + outros_nv)

        # Criar arquivo temporário
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmpfile:
            doc.save(tmpfile.name)
            file_path = tmpfile.name

        # Enviar e-mail
        try:
            msg = EmailMessage()
            msg['Subject'] = 'Nova Avaliação Neuropsicológica - Adulto'
            msg['From'] = st.secrets["email"]["from"]
            msg['To'] = "luan.gama.psicologo@gmail.com"  # ajuste aqui se necessário
            msg.set_content(f"Formulário preenchido por {nome} em {data_avaliacao}.")

            with open(file_path, 'rb') as f:
                msg.add_attachment(f.read(), maintype='application',
                                   subtype='vnd.openxmlformats-officedocument.wordprocessingml.document',
                                   filename=os.path.basename(file_path))

            with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
                smtp.login(st.secrets["email"]["from"], st.secrets["email"]["password"])
                smtp.send_message(msg)

            st.success("Formulário enviado com sucesso!")
        except Exception as e:
            st.error(f"Erro ao enviar: {e}")
        finally:
            os.remove(file_path)